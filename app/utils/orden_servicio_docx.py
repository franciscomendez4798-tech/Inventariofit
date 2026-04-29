"""
Genera el documento Word de la Orden de Servicio de Mantenimiento (R-AP-33-05-01).
Los logos institucionales van en el HEADER del documento (se repiten en cada página).
"""
import base64
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_LOGOS_DIR = os.path.join(os.path.dirname(__file__), 'logos')
_UAT_IMG  = os.path.join(_LOGOS_DIR, 'uat.jpg')
_SIGC_IMG = os.path.join(_LOGOS_DIR, 'sigc.png')
_FI_IMG   = os.path.join(_LOGOS_DIR, 'fi.jpg')


# ── Helpers ────────────────────────────────────────────────────────────────

def _no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def _checkbox(checked: bool) -> str:
    return '☑' if checked else '☐'


def _cell_border_none(cell):
    """Elimina todos los bordes de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_col_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ── Header del documento Word (se repite en cada página) ──────────────────

def _build_doc_header(doc):
    """
    Construye el encabezado oficial en el header del Word:
      ┌──────────┬──────────┬──────────┐
      │  UAT     │   SIGC   │    FI    │
      ├──────────┴──────────┴──────────┤
      │  ORDEN DE SERVICIO…  R-AP-…   │
      └─────────────────────────────  ┘
    """
    section = doc.sections[0]
    header  = section.header
    # Limpiar párrafo vacío por defecto
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # Fila 1: 3 logos
    tbl = header.add_table(rows=1, cols=3, width=Inches(5.4))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = [Inches(2.0), Inches(1.6), Inches(1.8)]
    logo_paths = [_UAT_IMG, _SIGC_IMG, _FI_IMG]
    logo_w     = [Inches(1.8), Inches(1.3), Inches(1.6)]

    for ci, (path, lw, cw) in enumerate(zip(logo_paths, logo_w, col_w)):
        cell = tbl.rows[0].cells[ci]
        _set_col_width(cell, cw / Inches(1))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(p)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if os.path.exists(path):
            run = p.add_run()
            run.add_picture(path, width=lw)
        else:
            p.add_run(os.path.basename(path)).font.size = Pt(8)

    # Fila 2: título del documento
    tbl2 = header.add_table(rows=1, cols=1, width=Inches(5.4))
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell2 = tbl2.rows[0].cells[0]
    _set_col_width(cell2, 5.4)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p2)
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after  = Pt(3)
    r1 = p2.add_run('ORDEN DE SERVICIO DE MANTENIMIENTO\n')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p2.add_run('R-AP-33-05-01')
    r2.font.size = Pt(10)

    # Ajustar top-margin para que el header tenga espacio suficiente
    section.header_distance = Cm(0.5)
    section.top_margin      = Cm(4.5)   # espacio para el header de ~3.5 cm
    section.bottom_margin   = Cm(2)
    section.left_margin     = Cm(2.5)
    section.right_margin    = Cm(2.5)


# ── Página 1: Solicitud ────────────────────────────────────────────────────

def _add_solicitud(doc, orden):
    fecha_sol = orden.fecha_solicitud.strftime('%d/%m/%Y') if orden.fecha_solicitud else '______________'
    servicios = orden.get_servicios()

    # Fecha (alineada a la derecha)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'Fecha de la solicitud: {fecha_sol}').font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    # Datos del solicitante
    for label, value in [
        ('Nombre del solicitante', orden.solicitante.nombre_completo if orden.solicitante else ''),
        ('Departamento o área',    orden.departamento.nombre if orden.departamento else ''),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_lbl = p.add_run(f'{label}: ')
        r_lbl.font.size = Pt(11)
        r_val = p.add_run(value)
        r_val.font.size = Pt(11)
        r_val.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Tipo de mantenimiento
    p = doc.add_paragraph()
    p.add_run('Tipo de Mantenimiento:').font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    for tipo in ['preventivo', 'correctivo']:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        p2.paragraph_format.space_after = Pt(1)
        p2.add_run(f'{_checkbox(orden.tipo_mantenimiento == tipo)}  {tipo.capitalize()}').font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Tipos de servicio (tabla 3×3)
    p = doc.add_paragraph()
    p.add_run('Tipo de servicio que solicita:').font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    grid = [
        [('plomeria','Plomería'),    ('pintura','Pintura'),        ('ac','A/C')],
        [('computo','Cómputo'),      ('jardineria','Jardinería'),   ('vehiculos','Vehículos')],
        [('electricidad','Electricidad'), ('limpieza','Limpieza'), None],
    ]
    tbl = doc.add_table(rows=3, cols=3)
    for ri, row_items in enumerate(grid):
        for ci, item in enumerate(row_items):
            cell = tbl.rows[ri].cells[ci]
            p_c = cell.paragraphs[0]
            _no_space(p_c)
            p_c.paragraph_format.space_before = Pt(1)
            p_c.paragraph_format.space_after  = Pt(1)
            if item:
                key, label = item
                p_c.add_run(f'{_checkbox(key in servicios)}  {label}').font.size = Pt(11)

    otro_val = orden.otro_servicio or ''
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.add_run(f'{_checkbox("otro" in servicios)}  Otro especifique: ').font.size = Pt(11)
    r2 = p.add_run(otro_val or '_' * 50)
    r2.font.size = Pt(11)
    r2.bold = bool(otro_val)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Descripción
    p = doc.add_paragraph()
    p.add_run('Descripción del servicio solicitado:').font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p2 = doc.add_paragraph()
    p2.add_run(orden.descripcion or '').font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    _add_caja_mantenimiento(doc, orden)


def _add_caja_mantenimiento(doc, orden):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run('Espacio para ser llenado únicamente por el personal que realizó el servicio')
    r.bold = True
    r.font.size = Pt(11)

    fecha_ej = orden.fecha_ejecucion.strftime('%d/%m/%Y') if orden.fecha_ejecucion else '_' * 15
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run(f'Fecha de ejecución del servicio: {fecha_ej}').font.size = Pt(11)

    if orden.servicio_realizado is None:
        r_txt = '☐ Sí   ☐ No'
    elif orden.servicio_realizado:
        r_txt = '☑ Sí   ☐ No'
    else:
        r_txt = '☐ Sí   ☑ No'
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    p3.add_run(f'¿Se realiza el servicio?: {r_txt}').font.size = Pt(11)

    motivo = orden.motivo_no_realizado or ''
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_after = Pt(4)
    p4.add_run('En caso no haberse realizado el servicio especificar motivo: ').font.size = Pt(11)
    r4b = p4.add_run(motivo if motivo else '_' * 70)
    r4b.font.size = Pt(11)
    r4b.bold = bool(motivo)


# ── Página 2: Firmas ───────────────────────────────────────────────────────

def _add_imagen_firma(cell, firma_b64: str, nombre: str, rol: str):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p)
    p.paragraph_format.space_before = Pt(4)

    if firma_b64:
        try:
            raw = base64.b64decode(firma_b64.split(',')[-1])
            buf = io.BytesIO(raw)
            run = p.add_run()
            run.add_picture(buf, width=Inches(1.6), height=Inches(0.6))
        except Exception:
            p.add_run('[ Firma ]').font.size = Pt(9)
    else:
        p.add_run('_' * 28).font.size = Pt(10)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p2)
    r2 = p2.add_run(nombre or '')
    r2.font.size = Pt(10)
    r2.bold = bool(nombre)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p3)
    p3.add_run('Nombre y Firma').font.size = Pt(9)

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(4)
    r4 = p4.add_run(rol)
    r4.font.size = Pt(9)
    r4.bold = True


def _add_firmas(doc, orden):
    doc.add_page_break()

    # Repetir caja de mantenimiento en página de firmas
    _add_caja_mantenimiento(doc, orden)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Tabla de firmas
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_data = [
        (orden.firma_realizo_b64,    orden.nombre_realizo or '',           'Realizó'),
        (orden.firma_superviso_b64,  orden.nombre_superviso or '',         'Supervisó'),
        (orden.firma_solicitante_b64,orden.nombre_solicitante_firma or '', 'Conformidad del solicitante'),
    ]
    for ci, (firma_b64, nombre, rol) in enumerate(col_data):
        cell = tbl.rows[0].cells[ci]
        cell.width = Inches(2.0)
        _add_imagen_firma(cell, firma_b64, nombre, rol)

    # Versión al pie
    doc.add_paragraph()
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ver = p_ver.add_run('Ver. 3    ACT. 22/04/2024')
    r_ver.font.size = Pt(9)
    r_ver.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Función principal ──────────────────────────────────────────────────────

def generar_docx_orden(orden) -> io.BytesIO:
    """
    Genera el .docx completo: header institucional (se repite) + solicitud + firmas.
    La encuesta de satisfacción se gestiona como notificación web, no en el documento.
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    _build_doc_header(doc)
    _add_solicitud(doc, orden)
    _add_firmas(doc, orden)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def guardar_docx_orden(orden, directorio: str) -> str:
    os.makedirs(directorio, exist_ok=True)
    nombre = f'OSM_{orden.folio}_{orden.fecha_solicitud.strftime("%Y%m%d")}.docx'
    ruta   = os.path.join(directorio, nombre)
    buf    = generar_docx_orden(orden)
    with open(ruta, 'wb') as f:
        f.write(buf.read())
    return ruta
