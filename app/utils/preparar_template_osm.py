"""
Script de uso único: toma el .docx oficial R-AP-33-05-01 y genera
template_osm.docx con etiquetas Jinja2 ({{...}}) en los campos
a rellenar. Solo conserva la primera página (elimina Encuesta).

Ejecución:
    venv/bin/python3 app/utils/preparar_template_osm.py
"""
import copy
import os
import sys

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ORIGINAL = "/Users/franmen/Downloads/R-AP-33-05-01 Orden de Servicio de Mantenimiento 2024 nvo logo.docx"
DESTINO  = os.path.join(os.path.dirname(__file__), "template_osm.docx")

# ── Helpers ────────────────────────────────────────────────────────────────

def _clear_paragraph(p):
    """Elimina todos los runs y drawings de un párrafo dejándolo vacío."""
    for child in list(p._element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'bookmarkStart', 'bookmarkEnd', 'del', 'ins',
                   'hyperlink', 'proofErr', 'permStart', 'permEnd'):
            p._element.remove(child)


def _add_run_with_tag(p, text, preserve_rpr=None):
    """Añade un run de texto plano, opcionalmente copiando el rPr del original."""
    r_el = etree.SubElement(p._element, qn('w:r'))
    if preserve_rpr is not None:
        rpr_copy = copy.deepcopy(preserve_rpr)
        r_el.insert(0, rpr_copy)
    t_el = etree.SubElement(r_el, qn('w:t'))
    t_el.text = text
    if text.startswith(' ') or text.endswith(' '):
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _get_first_rpr(p):
    """Devuelve el primer elemento rPr de los runs del párrafo, o None."""
    for r in p._element.findall(qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is not None:
            return rpr
    return None


def _remove_drawing_runs(p):
    """Elimina los runs que contienen drawings (imágenes inline de checkboxes)."""
    to_remove = []
    for child in p._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            if child.find(qn('w:drawing')) is not None:
                to_remove.append(child)
            else:
                # Run con texto vacío residual del drawing
                t = child.find(qn('w:t'))
                if t is None or (t.text or '').strip() == '':
                    to_remove.append(child)
    for el in to_remove:
        p._element.remove(el)


def _remove_numpPr(p):
    """Elimina la numeración (bullet de lista) de un párrafo."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)


def _delete_paragraphs_from(doc, start_idx):
    """Elimina todos los párrafos desde start_idx hasta el final del body."""
    body = doc.element.body
    # Recopilar todos los <w:p> directos del body
    paras = body.findall(qn('w:p'))
    for p_el in paras[start_idx:]:
        body.remove(p_el)


def _replace_text_in_run(p, old_substr, new_text):
    """Reemplaza texto dentro de los runs de un párrafo."""
    # Consolidar texto total del párrafo en runs
    for r in p._element.findall(qn('w:r')):
        t = r.find(qn('w:t'))
        if t is not None and t.text and old_substr in t.text:
            t.text = t.text.replace(old_substr, new_text)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


# ── Preparación ────────────────────────────────────────────────────────────

def preparar_template():
    doc = Document(ORIGINAL)
    paras = doc.paragraphs

    # ── 1. FECHA [01] ─────────────────────────────────────────────────────
    p01 = paras[1]
    _clear_paragraph(p01)
    _add_run_with_tag(p01, 'Fecha de la solicitud: {{fecha_solicitud}}')

    # ── 2. NOMBRE SOLICITANTE [03] ────────────────────────────────────────
    p03 = paras[3]
    _clear_paragraph(p03)
    _add_run_with_tag(p03, 'Nombre del solicitante: {{nombre_solicitante}}')

    # ── 3. DEPARTAMENTO [04] ──────────────────────────────────────────────
    p04 = paras[4]
    _clear_paragraph(p04)
    _add_run_with_tag(p04, 'Departamento o área: {{departamento}}')

    # ── 4. CHECKBOX PREVENTIVO [07] ───────────────────────────────────────
    p07 = paras[7]
    rpr = _get_first_rpr(p07)
    _remove_drawing_runs(p07)
    # Reemplazar runs de texto restante
    _clear_paragraph(p07)
    _add_run_with_tag(p07, '{{cb_preventivo}}  Preventivo', preserve_rpr=rpr)

    # ── 5. CHECKBOX CORRECTIVO [08] ───────────────────────────────────────
    p08 = paras[8]
    rpr = _get_first_rpr(p08)
    _remove_drawing_runs(p08)
    _clear_paragraph(p08)
    _add_run_with_tag(p08, '{{cb_correctivo}}  Correctivo', preserve_rpr=rpr)

    # ── 6. CHECKBOXES DE TIPO DE SERVICIO [13-20] ─────────────────────────
    servicios = [
        (13, 'cb_plomeria',     'Plomería'),
        (14, 'cb_computo',      'Cómputo'),
        (15, 'cb_electricidad', 'Electricidad'),
        (16, 'cb_pintura',      'Pintura'),
        (17, 'cb_jardineria',   'Jardinería'),
        (18, 'cb_limpieza',     'Limpieza'),
        (19, 'cb_ac',           'A/C'),
        (20, 'cb_vehiculos',    'Vehículos'),
    ]
    for idx, tag, label in servicios:
        p = paras[idx]
        rpr = _get_first_rpr(p)
        _remove_numpPr(p)
        _clear_paragraph(p)
        _add_run_with_tag(p, f'{{{{{tag}}}}}  {label}', preserve_rpr=rpr)

    # ── 7. OTRO ESPECIFIQUE [21] ──────────────────────────────────────────
    p21 = paras[21]
    _clear_paragraph(p21)
    _add_run_with_tag(p21, '{{cb_otro}}  Otro, especifique: {{otro_servicio}}')

    # ── 8. DESCRIPCIÓN [24] ───────────────────────────────────────────────
    p24 = paras[24]
    rpr = _get_first_rpr(p24)
    _clear_paragraph(p24)
    _add_run_with_tag(p24, '{{descripcion}}', preserve_rpr=rpr)

    # ── 9. ELIMINAR PÁGINAS 2 Y 3 (paras [25] en adelante) ───────────────
    # Dejar solo unos pocos párrafos en blanco al final para el espacio
    # del área de ejecución (llenado físico)
    _delete_paragraphs_from(doc, 25)

    doc.save(DESTINO)
    print(f"Template guardado en: {DESTINO}")
    print(f"Párrafos finales: {len(doc.paragraphs)}")


if __name__ == '__main__':
    preparar_template()
